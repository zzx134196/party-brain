"""文档解析模块 - 将PDF/Word文件解析为纯文本"""
import os
from loguru import logger


def parse_document(file_path: str) -> str:
    """解析文档，返回纯文本内容"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = file_path.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        return parse_pdf(file_path)
    elif ext in ("docx", "doc"):
        return parse_docx(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def parse_pdf(file_path: str) -> str:
    """解析PDF文件"""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        page_count = len(doc)
        texts = []
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            # 清洗：去除页眉页脚（简单策略：去除每页前后各2行中的短行）
            lines = text.strip().split("\n")
            cleaned_lines = clean_header_footer(lines)
            texts.append("\n".join(cleaned_lines))
        doc.close()

        full_text = "\n\n".join(texts)
        logger.info(f"PDF解析完成: {file_path}, 共{page_count}页, {len(full_text)}字")
        return full_text
    except ImportError:
        logger.error("PyMuPDF未安装，请运行: pip install PyMuPDF")
        raise


def parse_docx(file_path: str) -> str:
    """解析Word文档"""
    try:
        from docx import Document

        doc = Document(file_path)
        texts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    texts.append(row_text)

        full_text = "\n".join(texts)
        logger.info(f"Word解析完成: {file_path}, {len(full_text)}字")
        return full_text
    except ImportError:
        logger.error("python-docx未安装，请运行: pip install python-docx")
        raise


def clean_header_footer(lines: list, threshold: int = 10) -> list:
    """简单的页眉页脚清洗（去除短行）"""
    if len(lines) <= 4:
        return lines
    # 去除前2行和后2行中长度小于threshold的行
    cleaned = []
    for i, line in enumerate(lines):
        if (i < 2 or i >= len(lines) - 2) and len(line.strip()) < threshold:
            continue
        cleaned.append(line)
    return cleaned
