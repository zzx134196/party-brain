"""文档导出API - Word/PDF/Excel生成与下载"""
import io
import os
import uuid
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from loguru import logger

from app.models.database import get_db
from app.models.user import User
from app.core.auth import get_current_user
from app.config import settings

router = APIRouter(prefix="/api/export", tags=["文档导出"])

# 确保 exports 目录存在
EXPORT_DIR = os.path.abspath(settings.EXPORT_DIR)
os.makedirs(EXPORT_DIR, exist_ok=True)


@router.get("/download/{filename}")
async def download_file(filename: str):
    """通用文件下载接口 — 前端直接用链接即可下载"""
    # 安全检查：防止路径遍历
    safe_name = os.path.basename(filename)
    filepath = os.path.join(EXPORT_DIR, safe_name)
    if not os.path.isfile(filepath):
        raise HTTPException(status_code=404, detail="文件不存在或已过期")

    # 根据后缀确定 media_type
    ext = os.path.splitext(safe_name)[1].lower()
    media_types = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pdf": "application/pdf",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }
    from urllib.parse import quote
    return FileResponse(
        filepath,
        media_type=media_types.get(ext, "application/octet-stream"),
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(safe_name)}"},
    )


def _build_word_doc(title: str, content: str):
    """构建 Word 文档对象并返回 BytesIO"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for paragraph in content.split('\n'):
        stripped = paragraph.strip()
        if not stripped:
            doc.add_paragraph('')
            continue
        if stripped.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、')):
            doc.add_heading(stripped, level=2)
        elif stripped.startswith(('（一）', '（二）', '（三）', '（四）', '（五）')):
            doc.add_heading(stripped, level=3)
        else:
            p = doc.add_paragraph(stripped)
            p.style.font.size = Pt(12)

    doc.add_paragraph('')
    footer = doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    footer.style.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _build_pdf_doc(title: str, content: str):
    """构建 PDF 文档并返回 BytesIO"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_CENTER

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=25*mm, bottomMargin=25*mm)
    styles = getSampleStyleSheet()

    font_name = 'Helvetica'
    font_candidates = [
        '/System/Library/Fonts/STHeiti Light.ttc',
        '/System/Library/Fonts/PingFang.ttc',
        '/Library/Fonts/Arial Unicode.ttf',
        '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
        '/usr/share/fonts/wqy-zenhei/wqy-zenhei.ttc',
        '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
        'C:/Windows/Fonts/simhei.ttf',
    ]
    for font_path in font_candidates:
        try:
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('CJKFont', font_path))
                font_name = 'CJKFont'
                break
        except Exception:
            continue

    title_style = ParagraphStyle('ChTitle', parent=styles['Title'], fontName=font_name, fontSize=18, alignment=TA_CENTER)
    body_style = ParagraphStyle('ChBody', parent=styles['Normal'], fontName=font_name, fontSize=12, leading=20)
    heading_style = ParagraphStyle('ChHeading', parent=styles['Heading2'], fontName=font_name, fontSize=14)

    story = [Paragraph(title, title_style), Spacer(1, 12)]
    for line in content.split('\n'):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
            continue
        if stripped.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、')):
            story.append(Paragraph(stripped, heading_style))
        else:
            story.append(Paragraph(stripped, body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer


def auto_save_document(title: str, content: str) -> dict:
    """自动将文档内容保存为 Word 文件，返回下载链接信息。

    供 Agent Workflow 调用，生成文档后自动保存文件并在回复中嵌入下载链接。

    Returns:
        {"word_link": "/api/export/download/xxx.docx",
         "word_name": "xxx.docx"}
    """
    short_id = uuid.uuid4().hex[:8]
    import re as _re
    safe_title = _re.sub(r'[^\w\u4e00-\u9fff\-]', '', title)[:30] or "document"
    date_str = datetime.now().strftime('%Y%m%d')

    word_name = f"{safe_title}_{date_str}_{short_id}.docx"

    results = {}

    # 保存 Word
    try:
        buf = _build_word_doc(title, content)
        word_path = os.path.join(EXPORT_DIR, word_name)
        with open(word_path, "wb") as f:
            f.write(buf.read())
        results["word_link"] = f"/api/export/download/{word_name}"
        results["word_name"] = word_name
        logger.info(f"自动保存Word: {word_path}")
    except Exception as e:
        logger.error(f"自动保存Word失败: {e}")

    return results


class ExportWordRequest(BaseModel):
    title: str
    content: str


class ExportExcelRequest(BaseModel):
    columns: list
    rows: list
    sheet_name: str = "查询结果"


@router.post("/word")
async def export_word(
    req: ExportWordRequest,
    current_user: User = Depends(get_current_user),
):
    """导出Word文档"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 标题
        title_para = doc.add_heading(req.title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 正文 - 按段落添加
        for paragraph in req.content.split('\n'):
            stripped = paragraph.strip()
            if not stripped:
                doc.add_paragraph('')
                continue

            # 识别标题层级
            if stripped.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、')):
                doc.add_heading(stripped, level=2)
            elif stripped.startswith(('（一）', '（二）', '（三）', '（四）', '（五）')):
                doc.add_heading(stripped, level=3)
            else:
                p = doc.add_paragraph(stripped)
                p.style.font.size = Pt(12)

        # 页脚
        doc.add_paragraph('')
        footer = doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
        footer.style.font.size = Pt(9)

        # 写入内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        from urllib.parse import quote
        filename = f"{req.title}_{datetime.now().strftime('%Y%m%d')}.docx"
        encoded_filename = quote(filename)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
        )
    except Exception as e:
        logger.error(f"Word导出失败: {e}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.post("/pdf")
async def export_pdf(
    req: ExportWordRequest,
    current_user: User = Depends(get_current_user),
):
    """导出PDF文档"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_CENTER

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=25*mm, bottomMargin=25*mm)

        styles = getSampleStyleSheet()
        # 尝试注册中文字体（兼容多平台）
        font_name = 'Helvetica'
        font_candidates = [
            # macOS
            '/System/Library/Fonts/STHeiti Light.ttc',
            '/System/Library/Fonts/PingFang.ttc',
            '/Library/Fonts/Arial Unicode.ttf',
            # Linux (常见中文字体路径)
            '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
            '/usr/share/fonts/wqy-zenhei/wqy-zenhei.ttc',
            '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/google-noto-cjk/NotoSansCJK-Regular.ttc',
            # Windows
            'C:/Windows/Fonts/simhei.ttf',
            'C:/Windows/Fonts/simsun.ttc',
        ]
        for font_path in font_candidates:
            try:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('CJKFont', font_path))
                    font_name = 'CJKFont'
                    break
            except Exception:
                continue

        title_style = ParagraphStyle('ChTitle', parent=styles['Title'], fontName=font_name, fontSize=18, alignment=TA_CENTER)
        body_style = ParagraphStyle('ChBody', parent=styles['Normal'], fontName=font_name, fontSize=12, leading=20)
        heading_style = ParagraphStyle('ChHeading', parent=styles['Heading2'], fontName=font_name, fontSize=14)

        story = []
        story.append(Paragraph(req.title, title_style))
        story.append(Spacer(1, 12))

        for line in req.content.split('\n'):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 6))
                continue
            if stripped.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、')):
                story.append(Paragraph(stripped, heading_style))
            else:
                story.append(Paragraph(stripped, body_style))

        doc.build(story)
        buffer.seek(0)

        from urllib.parse import quote
        filename = f"{req.title}_{datetime.now().strftime('%Y%m%d')}.pdf"
        encoded_filename = quote(filename)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
        )
    except Exception as e:
        logger.error(f"PDF导出失败: {e}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.post("/excel")
async def export_excel(
    req: ExportExcelRequest,
    current_user: User = Depends(get_current_user),
):
    """导出Excel文件"""
    try:
        import pandas as pd

        df = pd.DataFrame(req.rows, columns=req.columns)
        buffer = io.BytesIO()
        df.to_excel(buffer, index=False, sheet_name=req.sheet_name, engine='openpyxl')
        buffer.seek(0)

        from urllib.parse import quote
        filename = f"{req.sheet_name}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        encoded_filename = quote(filename)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
        )
    except Exception as e:
        logger.error(f"Excel导出失败: {e}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")
