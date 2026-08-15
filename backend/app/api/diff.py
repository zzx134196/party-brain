"""文件差异对比API"""
import json
import os
import shutil
import tempfile
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from loguru import logger

from app.models.database import get_db
from app.models.user import User
from app.core.auth import get_current_user
from app.core.diff_engine import analyze_diff_with_llm, compute_similarity
from app.config import settings

router = APIRouter(prefix="/api/diff", tags=["文件差异对比"])


def _sse(payload: dict) -> str:
    """构造 SSE 数据帧"""
    return f"data: {json.dumps(payload, ensure_ascii=False, default=str)}\n\n"


def parse_uploaded_file(file_path: str) -> str:
    """解析上传的文件为纯文本（自包含，不依赖外部 parser 模块）"""
    ext = file_path.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            texts = [page.get_text("text") for page in doc]
            doc.close()
            return "\n\n".join(t.strip() for t in texts if t.strip())
        except ImportError:
            raise ValueError("PDF解析需要 PyMuPDF 库，请安装: pip install PyMuPDF")
    elif ext in ("docx", "doc"):
        try:
            from docx import Document
            doc = Document(file_path)
            texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip(" |"):
                        texts.append(row_text)
            return "\n".join(texts)
        except ImportError:
            raise ValueError("Word解析需要 python-docx 库，请安装: pip install python-docx")
    elif ext == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


@router.post("/compare")
async def compare_files(
    file1: UploadFile = File(...),
    file2: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    """上传两份文件进行差异对比"""
    allowed_exts = (".pdf", ".docx", ".doc", ".txt")
    for f in [file1, file2]:
        if not any(f.filename.lower().endswith(ext) for ext in allowed_exts):
            raise HTTPException(status_code=400, detail=f"不支持的文件格式: {f.filename}，仅支持PDF/Word/TXT")

    tmp_dir = tempfile.mkdtemp()
    try:
        # 保存临时文件
        path1 = os.path.join(tmp_dir, file1.filename)
        path2 = os.path.join(tmp_dir, file2.filename)

        content1 = await file1.read()
        content2 = await file2.read()

        with open(path1, "wb") as f:
            f.write(content1)
        with open(path2, "wb") as f:
            f.write(content2)

        # 解析文件
        text1 = parse_uploaded_file(path1)
        text2 = parse_uploaded_file(path2)

        if not text1.strip():
            raise HTTPException(status_code=400, detail=f"文件1 ({file1.filename}) 解析结果为空")
        if not text2.strip():
            raise HTTPException(status_code=400, detail=f"文件2 ({file2.filename}) 解析结果为空")

        # 执行差异分析
        result = await analyze_diff_with_llm(
            file1_name=file1.filename,
            file2_name=file2.filename,
            text1=text1,
            text2=text2,
        )

        result["file1"] = file1.filename
        result["file2"] = file2.filename

        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文件对比失败: {e}")
        raise HTTPException(status_code=500, detail=f"文件对比失败: {str(e)}")
    finally:
        # 清理临时文件
        shutil.rmtree(tmp_dir, ignore_errors=True)


@router.post("/compare/stream")
async def compare_files_stream(
    file1: UploadFile = File(...),
    file2: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    """上传两份文件进行差异对比（SSE 流式，前端可实时看到处理过程）"""
    allowed_exts = (".pdf", ".docx", ".doc", ".txt")
    for f in [file1, file2]:
        if not any(f.filename.lower().endswith(ext) for ext in allowed_exts):
            raise HTTPException(status_code=400, detail=f"不支持的文件格式: {f.filename}，仅支持PDF/Word/TXT")

    tmp_dir = tempfile.mkdtemp()

    async def generate():
        try:
            # 1. 解析文件1
            yield _sse({"type": "thinking_content", "text": "正在解析文件1...\n"})
            path1 = os.path.join(tmp_dir, file1.filename)
            with open(path1, "wb") as f:
                f.write(await file1.read())
            text1 = parse_uploaded_file(path1)
            if not text1.strip():
                raise HTTPException(status_code=400, detail=f"文件1 ({file1.filename}) 解析结果为空")
            yield _sse({"type": "thinking_content", "text": f"文件1解析完成（{len(text1)}字）。\n正在解析文件2...\n"})

            # 2. 解析文件2
            path2 = os.path.join(tmp_dir, file2.filename)
            with open(path2, "wb") as f:
                f.write(await file2.read())
            text2 = parse_uploaded_file(path2)
            if not text2.strip():
                raise HTTPException(status_code=400, detail=f"文件2 ({file2.filename}) 解析结果为空")
            yield _sse({"type": "thinking_content", "text": f"文件2解析完成（{len(text2)}字）。\n正在计算文本差异...\n"})

            # 3. 通知前端正在执行对比工具
            yield _sse({
                "type": "tool_calling",
                "tool": "file_diff",
                "args": {"file1": file1.filename, "file2": file2.filename},
            })

            # 4. 执行差异分析
            result = await analyze_diff_with_llm(
                file1_name=file1.filename,
                file2_name=file2.filename,
                text1=text1,
                text2=text2,
            )
            result["file1"] = file1.filename
            result["file2"] = file2.filename

            total = result.get("total_diffs", 0)
            similarity = result.get("similarity", 0)
            if result.get("fallback"):
                yield _sse({"type": "thinking_content", "text": "⚠️ AI语义差异分析暂不可用，已使用基础文本差异规则输出清单，建议人工复核。\n"})
            yield _sse({"type": "thinking_content", "text": f"对比完成，共发现 {total} 处差异，相似度 {similarity}%。\n"})

            yield _sse({
                "type": "tool_result",
                "tool": "file_diff",
                "success": True,
                "summary": f"对比完成，发现{total}处差异，相似度{similarity}%",
                "structured": {"type": "diff_report", "report": result},
            })
            yield _sse({
                "type": "done",
                "data": {"type": "diff_report", "report": result},
            })

        except HTTPException as e:
            detail = e.detail
            yield _sse({"type": "content", "text": f"⚠️ 文件差异对比失败：{detail}\n\n请稍后重试，或检查文件内容是否有效。"})
            yield _sse({"type": "done", "data": {"type": "diff_error", "error": detail}})
        except Exception as e:
            logger.error(f"文件差异对比流式处理失败: {e}")
            detail = str(e)
            yield _sse({"type": "content", "text": f"⚠️ 文件差异对比失败：{detail}\n\n请稍后重试；若多次失败，请联系管理员检查文件解析服务或AI模型服务。"})
            yield _sse({"type": "done", "data": {"type": "diff_error", "error": detail}})
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


@router.post("/compare-text")
async def compare_text(
    text1: str = Form(...),
    text2: str = Form(...),
    name1: str = Form("文件1"),
    name2: str = Form("文件2"),
    current_user: User = Depends(get_current_user),
):
    """直接对比两段文本"""
    if not text1.strip() or not text2.strip():
        raise HTTPException(status_code=400, detail="两段文本均不能为空")

    result = await analyze_diff_with_llm(
        file1_name=name1,
        file2_name=name2,
        text1=text1,
        text2=text2,
    )
    result["file1"] = name1
    result["file2"] = name2
    return result
