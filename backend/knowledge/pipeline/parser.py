"""文档解析器 - 支持PDF、Word文档（.docx/.doc）"""
import os
import subprocess
import tempfile
from pathlib import Path
from typing import Optional


def parse_document(file_path: str) -> str:
    """
    解析文档为纯文本
    
    Args:
        file_path: 文档路径
        
    Returns:
        解析后的纯文本内容
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf':
        return _parse_pdf(file_path)
    elif file_ext == '.docx':
        return _parse_docx(file_path)
    elif file_ext in ['.doc', '.wps']:
        return _parse_doc(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {file_ext}")


def _parse_pdf(file_path: str) -> str:
    """解析PDF文件 - 优先PyMuPDF，回退PyPDF2"""
    # 优先使用PyMuPDF（fitz），解析质量更高
    try:
        import fitz
        text = []
        with fitz.open(file_path) as doc:
            for page in doc:
                page_text = page.get_text()
                if page_text.strip():
                    text.append(page_text.strip())
        if text:
            return '\n\n'.join(text)
    except ImportError:
        pass
    except Exception:
        pass

    # 回退到PyPDF2
    try:
        import PyPDF2
        text = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return '\n\n'.join(text)
    except ImportError:
        pass
    except Exception:
        pass

    raise Exception(f"PDF解析失败: 需要安装 PyMuPDF 或 PyPDF2")


def _parse_docx(file_path: str) -> str:
    """解析.docx文件（新版Word格式）"""
    try:
        from docx import Document
        doc = Document(file_path)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return '\n\n'.join(text_parts)
    except ImportError:
        raise ImportError("需要安装python-docx: pip install python-docx")
    except Exception as e:
        raise Exception(f"Word文档解析失败: {e}")


def _parse_doc(file_path: str) -> str:
    """解析旧版.doc文件 - 使用macOS textutil或LibreOffice转换"""
    # 方案1: macOS自带textutil转txt
    try:
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
            tmp_path = tmp.name
        
        result = subprocess.run(
            ['textutil', '-convert', 'txt', '-output', tmp_path, file_path],
            capture_output=True, text=True, timeout=30
        )
        
        if result.returncode == 0 and os.path.exists(tmp_path):
            with open(tmp_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read().strip()
            os.unlink(tmp_path)
            if text:
                return text
        else:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    except FileNotFoundError:
        pass  # textutil不存在，跳过
    except Exception:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

    # 方案2: macOS textutil先转docx，再用python-docx解析
    try:
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_docx = tmp.name

        result = subprocess.run(
            ['textutil', '-convert', 'docx', '-output', tmp_docx, file_path],
            capture_output=True, text=True, timeout=30
        )

        if result.returncode == 0 and os.path.exists(tmp_docx):
            text = _parse_docx(tmp_docx)
            os.unlink(tmp_docx)
            if text:
                return text
        else:
            if os.path.exists(tmp_docx):
                os.unlink(tmp_docx)
    except FileNotFoundError:
        pass
    except Exception:
        if os.path.exists(tmp_docx):
            os.unlink(tmp_docx)

    # 方案3: LibreOffice转换（Linux服务器常用）
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'txt:Text', '--outdir', tmpdir, file_path],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0:
                txt_files = list(Path(tmpdir).glob('*.txt'))
                if txt_files:
                    with open(txt_files[0], 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read().strip()
                    if text:
                        return text
    except FileNotFoundError:
        pass
    except Exception:
        pass

    raise Exception(f"旧版.doc解析失败: 无法找到可用的转换工具(textutil/libreoffice)")
